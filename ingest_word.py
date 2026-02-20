
import os
import sys
import uuid
import logging
import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import pandas as pd
import psycopg2

# Import shared utilities
try:
    from ingest_common import get_db_connection, get_embeddings_batch
    from ingest_text import chunk_text_semantic, extract_financial_entities
except ImportError:
    logging.error("Could not import from ingest_common.py or ingest_text.py")
    sys.exit(1)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def extract_text_from_docx(file_path):
    """
    Extract text content from Word document.
    Returns full text as string.
    """
    try:
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        return '\n\n'.join(paragraphs)
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        raise e

def extract_tables_from_docx(file_path):
    """
    Extract tables from Word document.
    Returns list of DataFrames.
    """
    tables_data = []
    
    try:
        doc = Document(file_path)
        
        for table_idx, table in enumerate(doc.tables):
            # Convert table to 2D list
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
            
            if not data:
                continue
            
            # Try to use first row as header
            try:
                df = pd.DataFrame(data[1:], columns=data[0])
                # Remove completely empty columns
                df = df.dropna(axis=1, how='all')
                # Remove completely empty rows
                df = df.dropna(axis=0, how='all')
                
                if not df.empty:
                    tables_data.append({
                        'index': table_idx,
                        'df': df
                    })
            except Exception as e:
                logging.warning(f"Could not parse table {table_idx}: {e}")
                continue
    
    except Exception as e:
        logging.error(f"Error extracting tables from {file_path}: {e}")
        raise e
    
    return tables_data

def normalize_column_name(name):
    """
    Sanitize column names for SQL.
    """
    import re
    name = str(name).lower()
    name = re.sub(r'[^a-z0-9]', '_', name)
    name = re.sub(r'_+', '_', name)
    name = name.strip('_')
    
    if not name or name[0].isdigit():
        name = f"col_{name}"
    
    return name

def generate_word_summary(text, tables_count, file_name, openai_key):
    """
    Generate AI summary for Word document.
    """
    if not openai_key:
        return None
    
    try:
        from openai import OpenAI
        
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=openai_key
        )
        
        sample = text[:2000]
        
        prompt = (
            f"Analyze this Word document named '{file_name}'.\n"
            f"It contains {tables_count} tables.\n\n"
            f"Text preview:\n{sample}\n\n"
            f"Return a JSON object with:\n"
            f"1. 'category': Document type (e.g., 'Financial Report', 'Contract', 'Invoice')\n"
            f"2. 'summary': A concise 2-3 sentence summary\n"
            f"3. 'keywords': 5-10 keywords as comma-separated string"
        )
        
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a document analyst. Output valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300,
            temperature=0.3
        )
        
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logging.error(f"Failed to generate Word summary: {e}")
        return None

def create_word_tables(conn):
    """
    Create Word document tables in database.
    """
    with conn.cursor() as cur:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
        
        # Ensure files_metadata exists
        cur.execute("""
            CREATE TABLE IF NOT EXISTS files_metadata (
                file_id UUID PRIMARY KEY,
                file_name TEXT,
                uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                num_sheets INTEGER,
                num_pages INTEGER,
                summary TEXT,
                keywords TEXT,
                summary_embedding vector(3072),
                keywords_embedding vector(3072)
            );
        """)
        
        # Word document chunks (for text content)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS word_documents (
                doc_id UUID PRIMARY KEY,
                file_id UUID REFERENCES files_metadata(file_id) ON DELETE CASCADE,
                chunk_index INTEGER,
                chunk_text TEXT,
                entities JSONB,
                embedding vector(3072)
            );
        """)
        
        # Ensure sheets_metadata exists (for tables extracted from Word)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS sheets_metadata (
                sheet_id UUID PRIMARY KEY,
                file_id UUID REFERENCES files_metadata(file_id),
                sheet_name TEXT,
                table_name TEXT,
                num_rows INTEGER,
                num_columns INTEGER,
                category TEXT,
                summary TEXT,
                keywords TEXT,
                summary_embedding vector(3072),
                keywords_embedding vector(3072),
                columns_metadata JSONB,
                data_domain TEXT,
                sensitivity_level TEXT
            );
        """)
    
    conn.commit()
    logging.info("Word document tables verified.")

def process_word_file(file_path, db_url, openai_key, data_domain="general", sensitivity_level="internal", original_filename=None):
    """
    Main logic to ingest Word document.
    
    Extracts both text content and tables.
    Text is chunked and stored like text files.
    Tables are stored like Excel sheets.
    
    Args:
        file_path: Path to Word file
        db_url: Database URL
        openai_key: OpenAI API key
        data_domain: Data domain for RBAC
        sensitivity_level: Sensitivity level for RBAC
        original_filename: Original filename
    
    Returns: "SUCCESS" or raises Exception
    """
    file_path = os.path.abspath(file_path)
    file_name = original_filename if original_filename else os.path.basename(file_path)
    file_id = uuid.uuid4()
    
    logging.info(f"Processing Word file: {file_name} (ID: {file_id})")
    
    conn = get_db_connection(db_url)
    
    try:
        # 1. Extract text and tables
        full_text = extract_text_from_docx(file_path)
        tables = extract_tables_from_docx(file_path)
        
        logging.info(f"Extracted {len(full_text)} chars of text and {len(tables)} tables")
        
        # 2. Generate file-level metadata
        ai_analysis = generate_word_summary(full_text, len(tables), file_name, openai_key)
        
        if ai_analysis:
            file_summary = ai_analysis.get('summary', f"Word document: {file_name}")
            file_keywords = ai_analysis.get('keywords', file_name)
        else:
            file_summary = f"Word document with {len(tables)} tables: {file_name}"
            file_keywords = file_name
        
        # 3. Chunk text content
        text_chunks = chunk_text_semantic(full_text) if full_text.strip() else []
        logging.info(f"Created {len(text_chunks)} text chunks")
        
        # 4. Extract entities from chunks
        chunk_data = []
        for idx, chunk_text in enumerate(text_chunks):
            entities = extract_financial_entities(chunk_text)
            chunk_data.append({
                'index': idx,
                'text': chunk_text,
                'entities': entities
            })
        
        # 5. Prepare embeddings batch
        embedding_texts = [file_summary, file_keywords]
        
        # Add text chunk embeddings
        for chunk in chunk_data:
            embedding_texts.append(chunk['text'])
        
        # Add table summaries for embedding
        table_summaries = []
        for tbl in tables:
            df = tbl['df']
            cols_str = ', '.join(df.columns.tolist()[:5])
            summary = f"Table {tbl['index']} with columns: {cols_str}"
            table_summaries.append(summary)
            embedding_texts.append(summary)
        
        # 6. Generate embeddings
        logging.info(f"Generating embeddings for {len(embedding_texts)} items...")
        embeddings = get_embeddings_batch(embedding_texts, openai_key)
        
        file_summary_emb = embeddings[0]
        file_keywords_emb = embeddings[1]
        
        # Split embeddings for chunks and tables
        chunk_embeddings = embeddings[2:2+len(chunk_data)]
        table_embeddings = embeddings[2+len(chunk_data):]
        
        # 7. Database transaction
        with conn.cursor() as cur:
            # Insert file metadata
            cur.execute("""
                INSERT INTO files_metadata 
                (file_id, file_name, num_sheets, summary, keywords, summary_embedding, keywords_embedding)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (str(file_id), file_name, len(tables), file_summary, file_keywords, file_summary_emb, file_keywords_emb))
            
            # Insert text chunks
            if chunk_data:
                chunk_insert = """
                    INSERT INTO word_documents 
                    (doc_id, file_id, chunk_index, chunk_text, entities, embedding)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """
                
                chunk_records = []
                for i, chunk in enumerate(chunk_data):
                    doc_id = str(uuid.uuid4())
                    chunk_records.append((
                        doc_id,
                        str(file_id),
                        chunk['index'],
                        chunk['text'],
                        json.dumps(chunk['entities']),
                        chunk_embeddings[i]
                    ))
                
                cur.executemany(chunk_insert, chunk_records)
            
            # Insert tables
            for i, tbl in enumerate(tables):
                df = tbl['df']
                
                # Normalize column names
                df.columns = [normalize_column_name(c) for c in df.columns]
                
                # Deduplicate column names
                seen = {}
                new_cols = []
                for c in df.columns:
                    if c in seen:
                        seen[c] += 1
                        new_cols.append(f"{c}_{seen[c]}")
                    else:
                        seen[c] = 0
                        new_cols.append(c)
                df.columns = new_cols
                
                # Create table
                short_id = str(file_id).replace('-', '')[:8]
                table_name = f"word_table_{short_id}_{tbl['index']}"
                
                # Build CREATE TABLE
                col_defs = []
                for col in df.columns:
                    dtype = str(df[col].dtype)
                    if 'int' in dtype:
                        pgtype = 'BIGINT'
                    elif 'float' in dtype:
                        pgtype = 'DOUBLE PRECISION'
                    else:
                        pgtype = 'TEXT'
                    col_defs.append(f'"{col}" {pgtype}')
                
                create_sql = f"CREATE TABLE {table_name} ({', '.join(col_defs)});"
                cur.execute(create_sql)
                
                # Insert data
                df_clean = df.where(pd.notnull(df), None)
                vals = [tuple(x) for x in df_clean.to_numpy()]
                cols_sql = ', '.join([f'"{c}"' for c in df.columns])
                placeholders = ', '.join(['%s'] * len(df.columns))
                insert_sql = f"INSERT INTO {table_name} ({cols_sql}) VALUES ({placeholders})"
                
                cur.executemany(insert_sql, vals)
                
                # Insert sheet metadata
                sheet_id = uuid.uuid4()
                cur.execute("""
                    INSERT INTO sheets_metadata 
                    (sheet_id, file_id, sheet_name, table_name, num_rows, num_columns, 
                     summary, keywords, summary_embedding, keywords_embedding, data_domain, sensitivity_level)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    str(sheet_id), str(file_id), f"Table_{tbl['index']}", table_name,
                    len(df), len(df.columns), table_summaries[i], file_keywords,
                    table_embeddings[i], file_keywords_emb, data_domain, sensitivity_level
                ))
        
        conn.commit()
        logging.info(f"Successfully ingested {file_name} with {len(chunk_data)} text chunks and {len(tables)} tables")
        return "SUCCESS"
        
    except Exception as e:
        conn.rollback()
        logging.error(f"Error processing Word file {file_name}: {e}")
        raise e
    finally:
        conn.close()

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python ingest_word.py <path_to_word_file>")
        sys.exit(1)
    
    from dotenv import load_dotenv
    load_dotenv()
    
    word_file = sys.argv[1]
    db_url = os.getenv("DATABASE_URL")
    openai_key = os.getenv("OPENROUTER_API_KEY")
    
    if not db_url:
        logging.error("DATABASE_URL not found in environment")
        sys.exit(1)
    
    # Init tables
    conn = get_db_connection(db_url)
    create_word_tables(conn)
    conn.close()
    
    process_word_file(word_file, db_url, openai_key)
